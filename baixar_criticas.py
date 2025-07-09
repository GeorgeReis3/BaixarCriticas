import requests
from bs4 import BeautifulSoup
import sys
import os
import getpass
import openpyxl
from docx import Document
import platform
import subprocess
import shutil

class MoodleUtils:
    @staticmethod
    def completar_url(base_url, url_parcial):
        return url_parcial if url_parcial.startswith("http") else f"{base_url}/{url_parcial.lstrip('/')}"

    @staticmethod
    def limpar_nome(texto):
        return texto.replace("/", "_").replace("\\", "_").strip()

class Curso:
    def __init__(self, nome, url, categoria):
        self.nome = nome
        self.url = url
        self.categoria = categoria
        self.pasta = MoodleUtils.limpar_nome(nome)

class MoodleExtractor:
    BASE_URL = "https://www.gitevirtual.fab.mil.br"
    PALAVRAS_INDESEJADAS = [
        "NIL.", "nil.", "Nil.", "NIL", "nil", "Nil", "<br />", "<br>",
        "Nada a acrescentar.", "Sem comentário.", "Não há.", "Nada a acrescentar",
        "sem comentários", "Nenhuma sugestão a acrescentar.",
        "Não tenho comentários à acrescentar.", "não houve.", "sem sugestão.",
        "nada a comentar", "Nada a registrar.", "Nada a declarar.", "-", ".",
        "Nada a sugerir", "não há", "Nada a relatar.", "Nada a dizer.",
        "NÃO HÁ.", "Nada a relatar, pois todos os pontos foram bem abordados.",
        "Não tenho nada a sugerir, tendo em vista que o conteúdo foi bastante satisfatório para o meu aprendizado."
    ]

    def __init__(self, username, password):
        self.username = username
        self.password = password
        self.session = requests.Session()

    def acessar_url(self, url):
        response = self.session.get(url)
        response.encoding = 'utf-8'
        return BeautifulSoup(response.text, "html.parser")

    def buscar_link(self, soup, titulo, classe=None):
        if classe:
            return soup.find("a", class_=classe, title=titulo, href=True)
        return soup.find("a", title=titulo, href=True)

    def limpar_planilha(self, caminho_excel):
        print(f"    - Limpando planilha: {caminho_excel}")
        wb = openpyxl.load_workbook(caminho_excel)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str) and cell.value.strip() in self.PALAVRAS_INDESEJADAS:
                        cell.value = None
        wb.save(caminho_excel)
        print(f"    - Planilha limpa salva.")

    def baixar_excel(self, url, pasta_destino, nome_arquivo, nome_curso):
        os.makedirs(pasta_destino, exist_ok=True)
        analise_soup = self.acessar_url(url)
        botao_excel = analise_soup.find("button", string=lambda s: s and "Exportar para o Excel" in s)
        if botao_excel:
            form = botao_excel.find_parent("form")
            if form and form.has_attr("action"):
                export_url = MoodleUtils.completar_url(self.BASE_URL, form['action'])
                payload = {inp.get("name"): inp.get("value", "") for inp in form.find_all("input") if inp.get("name")}
                export_response = self.session.post(export_url, data=payload)
                if export_response.status_code == 200:
                    caminho = os.path.join(pasta_destino, f"{nome_arquivo}.xlsx")
                    with open(caminho, "wb") as f:
                        f.write(export_response.content)
                    print(f"    - Arquivo Excel salvo em '{caminho}'")
                    self.limpar_planilha(caminho)
                    self.gerar_docx_pdf(caminho, nome_curso, nome_arquivo)
                else:
                    print(f"    - Erro ao exportar: status {export_response.status_code}")
            else:
                print("    - [AVISO] Formulário de exportação não encontrado.")
        else:
            print("    - [AVISO] Botão 'Exportar para o Excel' não encontrado.")

    def gerar_docx_pdf(self, caminho_excel, nome_curso, nome_arquivo):
        modelo_path = os.path.join(os.path.dirname(__file__), "Modelo para crítica - (NÃO APAGAR).docx")
        pasta_destino = os.path.dirname(caminho_excel)
        caminho_docx = os.path.join(pasta_destino, f"{nome_arquivo}.docx")

        shutil.copy(modelo_path, caminho_docx)
        doc = Document(caminho_docx)

        doc.add_heading(f"Análise Crítica - {nome_curso}", level=1)
        doc.add_heading(nome_arquivo.replace("_", " "), level=2)

        wb = openpyxl.load_workbook(caminho_excel)
        sheet = wb.active

        linha = 6
        while linha <= sheet.max_row:
            pergunta = sheet[f"B{linha}"].value
            if pergunta:
                doc.add_paragraph(pergunta, style='Heading 3')

                cel_c = sheet[f"C{linha}"].value
                cel_d = sheet[f"D{linha}"].value

                if cel_c and cel_d:
                    alternativas = []
                    linha_qtd = linha + 1
                    col = 3
                    while sheet.cell(row=linha, column=col).value:
                        alt = sheet.cell(row=linha, column=col).value
                        qtd = sheet.cell(row=linha_qtd, column=col).value
                        if isinstance(qtd, (int, float)) and qtd > 0:
                            alternativas.append((alt, int(qtd)))
                        col += 1

                    total = sum(q[1] for q in alternativas)
                    if total > 0:
                        tabela = doc.add_table(rows=1, cols=3)
                        tabela.style = 'Table Grid'
                        hdr = tabela.rows[0].cells
                        hdr[0].text = 'Alternativa'
                        hdr[1].text = 'Quantidade'
                        hdr[2].text = 'Percentual (%)'

                        for alt, qtd in alternativas:
                            row = tabela.add_row().cells
                            row[0].text = str(alt)
                            row[1].text = str(qtd)
                            row[2].text = f"{round((qtd / total) * 100, 1)}%"
                    linha += 2
                else:
                    linha += 1
                    while linha <= sheet.max_row and not sheet[f"B{linha}"].value:
                        resposta = sheet[f"C{linha}"].value
                        if resposta and resposta.strip() and resposta.strip() not in self.PALAVRAS_INDESEJADAS:
                            doc.add_paragraph(f"- {resposta.strip()}")
                        linha += 1
            else:
                linha += 1

        doc.save(caminho_docx)
        print(f"    - Documento Word gerado com base no modelo: {caminho_docx}")

        if platform.system() == "Windows":
            try:
                import comtypes.client
                caminho_pdf = os.path.join(pasta_destino, f"{nome_arquivo}.pdf")
                word = comtypes.client.CreateObject('Word.Application')
                word.Visible = False
                docx_obj = word.Documents.Open(caminho_docx)
                docx_obj.SaveAs(caminho_pdf, FileFormat=17)
                docx_obj.Close()
                word.Quit()
                print(f"    - PDF gerado: {caminho_pdf}")
            except Exception as e:
                print(f"    - [ERRO] Falha ao gerar PDF: {e}")

    def login(self):
        print("[1] Fazendo login...")
        login_soup = self.acessar_url(f"{self.BASE_URL}/login/index.php")
        token_input = login_soup.find("input", {"name": "logintoken"})
        if not token_input or not token_input.has_attr('value'):
            print("[ERRO] Não foi possível encontrar o campo 'logintoken'")
            sys.exit(1)
        payload = {'username': self.username, 'password': self.password, 'logintoken': token_input['value']}
        login_response = self.session.post(f"{self.BASE_URL}/login/index.php", data=payload)
        if "login" in login_response.url or "senha incorreta" in login_response.text.lower():
            print("[ERRO] Falha no login.")
            sys.exit(1)
        print("[2] Login realizado com sucesso.")

    def listar_cursos_categorias(self):
        print("[3] Acessando categorias de cursos...")
        pagina_cursos = self.acessar_url(f"{self.BASE_URL}/course/index.php")
        categorias_alvo = ["Cursos EAD", "Cursos Presenciais", "CPT"]

        todos_cursos = []

        for categoria in categorias_alvo:
            link_categoria = pagina_cursos.find("a", string=lambda s: s and categoria in s)
            if link_categoria:
                url_categoria = MoodleUtils.completar_url(self.BASE_URL, link_categoria['href'])
                soup_categoria = self.acessar_url(url_categoria)
                cursos = soup_categoria.find_all("a", class_="aalink", href=True)
                for curso_link in cursos:
                    nome_curso = curso_link.get_text(strip=True)
                    url_curso = MoodleUtils.completar_url(self.BASE_URL, curso_link['href'])
                    todos_cursos.append(Curso(nome_curso, url_curso, categoria))

        if not todos_cursos:
            print("[ERRO] Nenhum curso encontrado.")
            sys.exit(1)

        print("\n[4] Lista de cursos:")
        numerados = []
        contador = 1
        for categoria in categorias_alvo:
            print(f"\n*** {categoria} ***")
            for curso in [c for c in todos_cursos if c.categoria == categoria]:
                print(f"[{contador}] {curso.nome}")
                numerados.append(curso)
                contador += 1

        print("""
Digite:
- os números dos cursos separados por espaço
- ou 'EAD' para todos de Cursos EAD
- ou 'PRES' para todos de Cursos Presenciais
- ou 'CPT' para todos de CPT
- ou ENTER para todos os cursos
""")

        selecao = input("Opção: ").strip().upper()

        cursos_selecionados = []
        if selecao == "EAD":
            cursos_selecionados = [c for c in numerados if c.categoria == "Cursos EAD"]
        elif selecao == "PRES":
            cursos_selecionados = [c for c in numerados if c.categoria == "Cursos Presenciais"]
        elif selecao == "CPT":
            cursos_selecionados = [c for c in numerados if c.categoria == "CPT"]
        elif selecao:
            indices = [int(i) for i in selecao.split() if i.isdigit() and 1 <= int(i) <= len(numerados)]
            cursos_selecionados = [numerados[i - 1] for i in indices]
        else:
            cursos_selecionados = numerados

        print(f"\n[5] {len(cursos_selecionados)} curso(s) selecionado(s).\n")
        return cursos_selecionados

    def processar_cursos(self, lista_cursos):
        for curso in lista_cursos:
            print(f"\n[6] Acessando curso: {curso.nome}")
            curso_soup = self.acessar_url(curso.url)

            critica_link = self.buscar_link(curso_soup, "Pesquisa", "dropdown-item") or self.buscar_link(curso_soup, "Pesquisa")

            if not critica_link:
                print("    - [AVISO] Link de 'Pesquisa' não encontrado.")
                continue

            print(f"    - Link de críticas encontrado: {critica_link['href']}")
            pesquisa_soup = self.acessar_url(critica_link['href'])
            tabela = pesquisa_soup.find("table", class_="generaltable")

            if not tabela:
                print("    - [AVISO] Tabela de críticas não encontrada.")
                continue

            linhas = tabela.find_all("tr")[1:]

            for linha in linhas:
                colunas = linha.find_all("td")
                if len(colunas) >= 2:
                    titulo_quest = colunas[1].get_text(strip=True)
                    nome_arquivo = MoodleUtils.limpar_nome(titulo_quest)
                    url_quest = MoodleUtils.completar_url(self.BASE_URL, colunas[1].find("a", href=True)['href'])
                    quest_soup = self.acessar_url(url_quest)
                    link_analise = quest_soup.find("a", title="Análise")
                    if link_analise:
                        print(f"    - Acessando link de Análise: {link_analise['href']}")
                        self.baixar_excel(MoodleUtils.completar_url(self.BASE_URL, link_analise['href']), curso.pasta, nome_arquivo, curso.nome)

    def encerrar_sessao(self):
        self.session.close()
        print("[FIM] Sessão encerrada.")

if __name__ == "__main__":
    print("==== Moodle Critic Report ====")
    username = input("Digite o usuário: ").strip()
    password = getpass.getpass("Digite a senha: ").strip()
    extractor = MoodleExtractor(username, password)
    extractor.login()
    cursos = extractor.listar_cursos_categorias()
    extractor.processar_cursos(cursos)
    extractor.encerrar_sessao()
