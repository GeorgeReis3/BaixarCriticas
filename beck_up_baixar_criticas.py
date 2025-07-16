import requests
from bs4 import BeautifulSoup
import sys
import os
import getpass
import openpyxl
from docx import Document
import platform
import subprocess
import tkinter as tk
from tkinter import messagebox, simpledialog

class MoodleUtils:
    @staticmethod
    def completar_url(base_url, url_parcial):
        return url_parcial if url_parcial.startswith("http") else f"{base_url}/{url_parcial.lstrip('/')}"

    @staticmethod
    def limpar_nome(texto):
        return texto.replace("/", "_").replace("\\", "_").strip()

class Curso:
    def __init__(self, nome, url, categoria):
        self.nome = nome
        self.url = url
        self.categoria = categoria
        self.pasta = MoodleUtils.limpar_nome(nome)

class MoodleExtractor:
    BASE_URL = "https://www.gitevirtual.fab.mil.br"
    PALAVRAS_INDESEJADAS = [
        "NIL.", "nil.", "Nil.", "NIL", "nil", "Nil", "<br />", "<br>",
        "Nada a acrescentar.", "Sem comentário.", "Não há.", "Nada a acrescentar",
        "sem comentários", "Nenhuma sugestão a acrescentar.",
        "Não tenho comentários à acrescentar.", "não houve.", "sem sugestão.",
        "nada a comentar", "Nada a registrar.", "Nada a declarar.", "-", ".",
        "Nada a sugerir", "não há", "Nada a relatar.", "Nada a dizer.",
        "NÃO HÁ.", "Nada a relatar, pois todos os pontos foram bem abordados.",
        "Não tenho nada a sugerir, tendo em vista que o conteúdo foi bastante satisfatório para o meu aprendizado."
    ]

    def __init__(self, username, password):
        self.username = username
        self.password = password
        self.session = requests.Session()

    def acessar_url(self, url):
        response = self.session.get(url)
        response.encoding = 'utf-8'
        return BeautifulSoup(response.text, "html.parser")

    def buscar_link(self, soup, titulo, classe=None):
        if classe:
            return soup.find("a", class_=classe, title=titulo, href=True)
        return soup.find("a", title=titulo, href=True)

    def baixar_excel(self, url, pasta_destino, nome_arquivo, nome_curso):
        os.makedirs(pasta_destino, exist_ok=True)
        analise_soup = self.acessar_url(url)
        botao_excel = analise_soup.find("button", string=lambda s: s and "Exportar para o Excel" in s)
        if botao_excel:
            form = botao_excel.find_parent("form")
            if form and form.has_attr("action"):
                export_url = MoodleUtils.completar_url(self.BASE_URL, form['action'])
                payload = {inp.get("name"): inp.get("value", "") for inp in form.find_all("input") if inp.get("name")}
                export_response = self.session.post(export_url, data=payload)
                if export_response.status_code == 200:
                    caminho = os.path.join(pasta_destino, f"{nome_arquivo}.xlsx")
                    with open(caminho, "wb") as f:
                        f.write(export_response.content)
                    print(f"    - Arquivo Excel salvo em '{caminho}'")
                    self.limpar_planilha(caminho)
                    self.gerar_docx_pdf(caminho, nome_curso, nome_arquivo)
                else:
                    print(f"    - Erro ao exportar: status {export_response.status_code}")
            else:
                print("    - [AVISO] Formulário de exportação não encontrado.")
        else:
            print("    - [AVISO] Botão 'Exportar para o Excel' não encontrado.")

    def limpar_planilha(self, caminho_arquivo):
        try:
            wb = openpyxl.load_workbook(caminho_arquivo)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            if cell.value.strip() in self.PALAVRAS_INDESEJADAS:
                                cell.value = ""
            wb.save(caminho_arquivo)
            print(f"    - Palavras indesejadas removidas de '{caminho_arquivo}'")
        except Exception as e:
            print(f"    - [ERRO] Falha ao limpar '{caminho_arquivo}': {e}")

    def gerar_docx_pdf(self, caminho_planilha, nome_curso, nome_arquivo):
        try:
            import re
            modelo_path = os.path.join(os.getcwd(), "Modelo para crítica - (NÃO APAGAR).docx")
            doc = Document(modelo_path)

            wb = openpyxl.load_workbook(caminho_planilha)
            sheet = wb.active

            titulo = sheet['A2'].value or ""
            n_respostas = "0"
            if titulo:
                match = re.search(r'(\d+)', titulo)
                if match:
                    n_respostas = match.group(1)

            for par in doc.paragraphs:
                if "{{CURSO}}" in par.text:
                    par.text = par.text.replace("{{CURSO}}", nome_curso)
                if "{{DISCIPLINA}}" in par.text:
                    par.text = par.text.replace("{{DISCIPLINA}}", nome_arquivo)
                if "{{NRESPOSTAS}}" in par.text:
                    par.text = par.text.replace("{{NRESPOSTAS}}", n_respostas)

            num_respostas = int(n_respostas) if n_respostas.isdigit() else 0

            for row_idx in range(6, sheet.max_row + 1):
                pergunta = sheet[f'B{row_idx}'].value
                if pergunta:
                    doc.add_paragraph("")
                    doc.add_paragraph(pergunta)
                    doc.add_paragraph("")

                    respostas = [cell.value for cell in sheet[row_idx][2:]]
                    porcentagens = [cell.value for cell in sheet[row_idx + 2][2:]] if row_idx + 2 <= sheet.max_row else []

                    if respostas and porcentagens and any(respostas):
                        tabela = doc.add_table(rows=0, cols=2)
                        tabela.style = 'Plain Table 4'
                        for resp, perc in zip(respostas, porcentagens):
                            if resp is not None and resp != "":
                                linha = tabela.add_row().cells
                                linha[0].text = str(resp)
                                linha[1].text = f"{(perc * 100):.2f}%" if isinstance(perc, (int, float)) else str(perc)
                    else:
                        tabela2 = doc.add_table(rows=0, cols=1)

                        tabela2.style = 'Plain Table 4'
                        for i in range(num_respostas):
                            comentario = sheet.cell(row=row_idx + i, column=3).value
                            if comentario and comentario != "":
                                linha = tabela2.add_row().cells
                                linha[0].text = str(comentario)

            docx_path = caminho_planilha.replace(".xlsx", ".docx")
            doc.save(docx_path)
            print(f"    - Documento salvo em '{docx_path}'")

            self.converter_para_pdf(docx_path)

        except Exception as e:
            print(f"    - [ERRO] Falha ao gerar DOCX/PDF: {e}")

    def converter_para_pdf(self, caminho_docx):
        try:
            sistema = platform.system()
            if sistema == "Windows":
                from docx2pdf import convert
                convert(caminho_docx)
                print(f"    - PDF gerado para '{caminho_docx.replace('.docx', '.pdf')}'")
            elif sistema == "Linux":
                subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", caminho_docx], check=True)
                print(f"    - PDF gerado para '{caminho_docx.replace('.docx', '.pdf')}'")
            else:
                print("    - [AVISO] Conversão para PDF não suportada neste sistema.")
        except Exception as e:
            print(f"    - [ERRO] Falha ao converter PDF: {e}")

    def login(self):
        print("[1] Fazendo login...")
        login_soup = self.acessar_url(f"{self.BASE_URL}/login/index.php")
        token_input = login_soup.find("input", {"name": "logintoken"})
        if not token_input or not token_input.has_attr('value'):

            raise Exception("[ERRO] Não foi possível encontrar o campo 'logintoken'")
        payload = {'username': self.username, 'password': self.password, 'logintoken': token_input['value']}
        login_response = self.session.post(f"{self.BASE_URL}/login/index.php", data=payload)
        if "login" in login_response.url or "senha incorreta" in login_response.text.lower():
            raise Exception("[ERRO] Falha no login.")
        print("[2] Login realizado com sucesso.")

    def encerrar_sessao(self):
        self.session.close()
        print("[FIM] Sessão encerrada.")

    def listar_cursos_categorias(self):
        print("[3] Acessando categorias de cursos...")
        pagina_cursos = self.acessar_url(f"{self.BASE_URL}/course/index.php")
        categorias_alvo = ["Cursos EAD", "Cursos Presenciais", "CPT"]


        todos_cursos = []

        for categoria in categorias_alvo:
            link_categoria = pagina_cursos.find("a", string=lambda s: s and categoria in s)
            if link_categoria:
                url_categoria = MoodleUtils.completar_url(self.BASE_URL, link_categoria['href'])
                soup_categoria = self.acessar_url(url_categoria)
                cursos = soup_categoria.find_all("a", class_="aalink", href=True)
                for curso_link in cursos:
                    nome_curso = curso_link.get_text(strip=True)
                    url_curso = MoodleUtils.completar_url(self.BASE_URL, curso_link['href'])
                    todos_cursos.append(Curso(nome_curso, url_curso, categoria))

        if not todos_cursos:
            print("[ERRO] Nenhum curso encontrado.")
            sys.exit(1)

        return todos_cursos

    def processar_cursos(self, lista_cursos):
        for curso in lista_cursos:
            print(f"\n[6] Acessando curso: {curso.nome}")
            curso_soup = self.acessar_url(curso.url)
            critica_link = self.buscar_link(curso_soup, "Pesquisa", "dropdown-item") or self.buscar_link(curso_soup, "Pesquisa")
            if not critica_link:
                print("    - [AVISO] Link de 'Pesquisa' não encontrado.")
                continue

            print(f"    - Link de críticas encontrado: {critica_link['href']}")
            pesquisa_soup = self.acessar_url(critica_link['href'])
            tabela = pesquisa_soup.find("table", class_="generaltable")
            if not tabela:
                print("    - [AVISO] Tabela de críticas não encontrada.")
                continue

            linhas = tabela.find_all("tr")[1:]
            for linha in linhas:
                colunas = linha.find_all("td")
                if len(colunas) >= 2:
                    titulo_quest = colunas[1].get_text(strip=True)
                    nome_arquivo = MoodleUtils.limpar_nome(titulo_quest)
                    url_quest = MoodleUtils.completar_url(self.BASE_URL, colunas[1].find("a", href=True)['href'])
                    quest_soup = self.acessar_url(url_quest)
                    link_analise = quest_soup.find("a", title="Análise")
                    if link_analise:
                        print(f"    - Acessando link de Análise: {link_analise['href']}")
                        self.baixar_excel(MoodleUtils.completar_url(self.BASE_URL, link_analise['href']), curso.pasta, nome_arquivo, curso.nome)


def iniciar_interface_grafica():
    def executar_geracao():
        username = entrada_usuario.get().strip()
        password = entrada_senha.get().strip()
        if not username or not password:
            messagebox.showerror("Erro", "Usuário e senha são obrigatórios.")
            return
        try:
            extractor = MoodleExtractor(username, password)
            extractor.login()
            janela.destroy()

            # Realiza a listagem e captura dos cursos
            print("[GUI] Listando cursos...")
            cursos = extractor.listar_cursos_categorias()

            # Interface de seleção de cursos com scroll e seleção múltipla
            nova_janela = tk.Tk()
            nova_janela.title("Seleção de Cursos")
            nova_janela.geometry("600x500")

            tk.Label(nova_janela, text="Selecione os cursos desejados:").pack(pady=5)

            frame_lista = tk.Frame(nova_janela)
            frame_lista.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

            scrollbar = tk.Scrollbar(frame_lista)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

            lista_cursos = tk.Listbox(frame_lista, selectmode=tk.MULTIPLE, yscrollcommand=scrollbar.set, width=80, height=20)
            lista_cursos.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

            scrollbar.config(command=lista_cursos.yview)

            cursos_numerados = []
            categorias = ["Cursos EAD", "Cursos Presenciais", "CPT"]
            for categoria in categorias:
                lista_cursos.insert(tk.END, f"--- {categoria} ---")
                for curso in [c for c in cursos if c.categoria == categoria]:
                    lista_cursos.insert(tk.END, curso.nome)
                    cursos_numerados.append(curso)

            def processar():
                selecionados_indices = lista_cursos.curselection()
                selecionados = []
                for i in selecionados_indices:
                    item = lista_cursos.get(i)
                    if item.startswith("---"):
                        continue  # Ignora divisores de categoria
                    # Encontra o curso correspondente na lista numerada
                    for curso in cursos:
                        if curso.nome == item:
                            selecionados.append(curso)
                            break

                if not selecionados:
                    messagebox.showerror("Erro", "Nenhum curso selecionado.")
                    return

                nova_janela.destroy()
                extractor.processar_cursos(selecionados)
                extractor.encerrar_sessao()
                messagebox.showinfo("Concluído", "Processamento finalizado com sucesso!")

            tk.Button(nova_janela, text="Processar Selecionados", command=processar).pack(pady=10)

            nova_janela.mainloop()

        except Exception as e:
            messagebox.showerror("Erro", str(e))

    janela = tk.Tk()
    janela.title("Gerador de Relatórios Moodle")
    janela.geometry("400x200")

    tk.Label(janela, text="Usuário:").pack(pady=5)
    entrada_usuario = tk.Entry(janela, width=30)
    entrada_usuario.pack(pady=5)

    tk.Label(janela, text="Senha:").pack(pady=5)
    entrada_senha = tk.Entry(janela, show="*", width=30)
    entrada_senha.pack(pady=5)

    tk.Button(janela, text="Executar", command=executar_geracao).pack(pady=20)

    janela.mainloop()

if __name__ == "__main__":
    iniciar_interface_grafica()


    # Falta alterar/formatar a largura da coluna para os números ficarem mais a direita e os textos em menos linhas
    # Falta verificar a possibilidade de transpor as tabelas, de forma a ocuparem menos espaço. (Pelo menos as tabelas de 0 a 10)
    # Falta formatar os parágrafos.
    # Falta formatar dentro da tabela (a primeira linha e a primeira coluna estão saindo em negrito)
    

